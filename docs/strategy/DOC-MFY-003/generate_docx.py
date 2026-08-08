"""Generate DOCX deliverables for DOC-MFY-003."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import datetime

OUT_DIR = Path(__file__).resolve().parent / "docx"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ── Styling helpers ────────────────────────────────────────────
HEADER_COLOR = "1B3A5C"
ACCENT_COLOR = "0D6EFD"
TEXT_COLOR = "212529"
BORDER_COLOR = "DEE2E6"


def set_cell_shading(cell, color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, HEADER_COLOR)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(*_hex_to_rgb(TEXT_COLOR))
            if c == 0:
                run.bold = True
            if r % 2 == 1:
                set_cell_shading(cell, "F8F9FA")

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table


def _hex_to_rgb(hex_str):
    h = hex_str.lstrip("#")
    return tuple(int(h[i : i + 2], 16) for i in (0, 2, 4))


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*_hex_to_rgb(HEADER_COLOR))
    return h


def add_meta_table(doc):
    """Add metadata table at top of document."""
    meta = [
        ("文档编号", "DOC-MFY-003"),
        ("版本代号", "v0.4 — Acoustic Compliance Upgrade (ACU)"),
        ("日期", "2026-07-02 15:37 (Asia/Singapore, UTC+8)"),
        ("公司", "荣景文川（深圳）科技有限公司"),
        ("上游文档", "DOC-MFY-002（声学理论合规审计与第二代研发路线图）"),
        ("下游输出", "EXP-MFY 声学实验任务 / ENG-MFY 工程实现任务"),
        ("状态", "草稿 — 待 Founder/CTO 签核"),
    ]
    add_styled_table(doc, ["属性", "内容"], meta, [3.5, 13])


def add_page_number(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("— ")
        run.font.size = Pt(8)
        # Page number field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run._r.append(instrText)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar2)
        run2 = p.add_run(" —")
        run2.font.size = Pt(8)


# ═══════════════════════════════════════════════════════════════
# 1. 项目章程合订版
# ═══════════════════════════════════════════════════════════════
def build_charter():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Cover page ──
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DOC-MFY-003")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(*_hex_to_rgb(HEADER_COLOR))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Moodify Acoustic Compliance Upgrade v0.4\n项目章程（合订版）")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(*_hex_to_rgb(ACCENT_COLOR))

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("荣景文川（深圳）科技有限公司\n2026年07月02日\n内部研发管理资产")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(*_hex_to_rgb("6C757D"))

    doc.add_page_break()

    # ── Chapter 1: 项目章程正文 ──
    add_heading(doc, "第一部分：项目章程正文", 1)
    add_meta_table(doc)

    add_heading(doc, "1. 项目定义", 2)
    doc.add_paragraph(
        "v0.4 (Acoustic Compliance Upgrade) 是 Moodify 的声学合规补完版本。"
        "其核心不是新增功能，而是修复 DOC-MFY-002 审计中发现的声学理论与实现缺口。"
    )

    add_heading(doc, "1.1 v0.4 战略定位", 3)
    strategy_data = [
        ["v0.1.0 (当前)", "v0.4 (本次)", "v0.5+ (后续)"],
        ["声学合规度: 65/100", "声学合规度: 88/100", "感知对齐度: 17/100+"],
        ["感知对齐度: 0/100", "感知对齐度: 5/100", "智能层扩展"],
        ["审计闭环: 6.0%", "审计闭环: 25%+", "闭环效率: 60%+"],
    ]
    add_styled_table(doc, strategy_data[0], strategy_data[1:], [4.5, 5, 5])

    add_heading(doc, "1.2 成功判定（6 条全部满足）", 3)
    success = [
        ["G1", "P0 缺陷 DEF-001 (RBJ EQ) 和 DEF-002 (Schroeder Reverb) 修复并验收通过"],
        ["G2", "AEP-ACU-001~010 中 P0 和 P1 任务的验收证据全部产出"],
        ["G3", "声学合规评分 A_compliance 从 65 → ≥ 85"],
        ["G4", "所有修复通过 ruff lint + pytest -m v01 + pytest 全量"],
        ["G5", "回归测试：20 首基线音频的处理前后 MRS 差值变化 < 2.0 分"],
        ["G6", "DOC-MFY-003 全部交付物归档完成"],
    ]
    for g_id, g_text in success:
        p = doc.add_paragraph()
        run = p.add_run(f"[{g_id}] ")
        run.bold = True
        p.add_run(g_text)

    add_heading(doc, "2. 上游审计发现映射", 2)
    mapping = [
        ["DEF-001 (FFT EQ → RBJ)", "AEP-ACU-002", "P0", "RBJ biquad EQ 替换"],
        ["DEF-002 (Schroeder 缺全通)", "AEP-ACU-001", "P0", "全通滤波器 + 双声道混响"],
        ["DEF-006 (频段不一致)", "AEP-ACU-004", "P1", "7 频段统一 + 5-8 kHz 区间"],
        ["DEF-001/NEW-04 (真峰值)", "AEP-ACU-005", "P1", "True Peak Limiter"],
        ["DEF-005 (mel/bark/chroma/F0)", "AEP-ACU-006", "P1", "Mel/Bark/ERB 感知尺度"],
        ["DEF-005 (F0)", "AEP-ACU-008", "P2", "F0/Pitch Stability"],
        ["DEF-005 (chroma)", "AEP-ACU-009", "P2", "Chroma/Key/Harmony"],
        ["DEF-004 (掩蔽模型)", "AEP-ACU-007", "P2", "心理声学掩蔽初版"],
        ["新增 (HPSS 残差)", "AEP-ACU-003", "P0", "HPSS 残差分量守恒审计"],
        ["新增 (MRS 鲁棒化)", "AEP-ACU-010", "P2", "MRS 参考集鲁棒化"],
    ]
    add_styled_table(doc, ["DOC-MFY-002 缺陷", "v0.4 任务", "优先级", "说明"], mapping, [4, 3, 1.5, 7])

    add_heading(doc, "3. 团队与角色", 2)
    roles = [
        ["Founder / CTO", "最终签核、优先级确认、IP 策略决策"],
        ["Claude A (交接官)", "huliye24 — 任务分发、跨文档一致性、封口验收"],
        ["Claude C (审计官)", "代码审查、合规验证、bug 检测（只查不修）"],
        ["DeepSeek (AI Worker)", "AEP 任务执行、实验设计、报告生成"],
        ["Codex (实现)", "ENG-MFY 工程实现"],
    ]
    add_styled_table(doc, ["角色", "职责"], roles, [4.5, 11])

    add_heading(doc, "4. 关键里程碑", 2)
    milestones = [
        ["M1", "项目章程签核", "2026-07-03", "Founder/CTO 批准 DOC-MFY-003"],
        ["M2", "P0 修复完成", "2026-07-05", "AEP-ACU-001/002/003 验收通过"],
        ["M3", "P1 补完完成", "2026-07-12", "AEP-ACU-004/005/006 验收通过"],
        ["M4", "回归测试通过", "2026-07-13", "全量 pytest + MRS 回归"],
        ["M5", "P2 合并判断", "2026-07-14", "决定 P2 合入或推迟"],
        ["M6", "v0.4 封口", "2026-07-15", "全部交付物归档 + git tag"],
    ]
    add_styled_table(doc, ["编号", "里程碑", "日期", "条件"], milestones, [1.2, 3.5, 2.5, 8])

    doc.add_page_break()

    # ── Chapter 2: Scope Boundary ──
    add_heading(doc, "第二部分：范围边界", 1)
    doc.add_paragraph(
        "原则：范围边界是项目章程的承重墙。v0.4 的失败模式不是「做得不够」，"
        "而是「做进去了不该做的事」。"
    )

    add_heading(doc, "In Scope — 必做项", 2)
    scope_in = [
        ["SCOPE-01", "RBJ biquad EQ 替换 FFT sigmoid/Gaussian EQ", "P0", "DEF-001"],
        ["SCOPE-02", "Schroeder 混响全通滤波器级补全", "P0", "DEF-002"],
        ["SCOPE-03", "HPSS 残差分量保留与能量守恒审计", "P0", "新增"],
        ["SCOPE-04", "7 频段默认启用并修复 5-8 kHz 空白", "P1", "DEF-006"],
        ["SCOPE-05", "True Peak Limiter (ITU-R BS.1771)", "P1", "NEW-04"],
        ["SCOPE-06", "Mel / Bark / ERB 感知频率尺度分析层", "P1", "DEF-005"],
        ["SCOPE-07", "心理声学掩蔽最小模型", "P2", "DEF-004"],
        ["SCOPE-08", "F0 / Pitch Stability 分析入口", "P2", "DEF-005"],
        ["SCOPE-09", "Chroma / Key / Harmony 分析入口", "P2", "DEF-005"],
        ["SCOPE-10", "MRS 参考集鲁棒化", "P2", "新增"],
    ]
    add_styled_table(doc, ["编号", "条目", "优先级", "来源"], scope_in, [2, 9.5, 1.5, 2.5])

    add_heading(doc, "Out of Scope — 绝不做", 2)
    scope_out = [
        ["OUT-01", "不在 v0.4 中训练任何深度学习模型", "硬件不足 + 非合规补完范围"],
        ["OUT-02", "不将 Demucs/Spleeter 等 DL 源分离集成到主处理路径", "HPSS 够用，留给 v0.6+"],
        ["OUT-03", "不在此版本接入新的外部 AI API", "不扩大外部依赖面"],
        ["OUT-04", "不做 UI/UX 改版", "v0.4 是声学引擎升级"],
        ["OUT-05", "不做商业发布承诺", "内部研发里程碑"],
        ["OUT-06", "不写对外营销材料", "内部研发管理资产"],
        ["OUT-07", "不做大规模架构重构", "在现有模块边界内完成"],
        ["OUT-08", "不做新的 CLI 命令或 API 端点", "对现有接口透明"],
        ["OUT-09", "不做 GPU 硬件采购或云迁移", "硬件决策独立于版本"],
    ]
    add_styled_table(doc, ["编号", "条目", "理由"], scope_out, [2, 8, 5.5])

    doc.add_page_break()

    # ── Chapter 3: AEP Task Cards ──
    add_heading(doc, "第三部分：AEP-ACU-001~010 任务卡", 1)
    aep_summary = [
        ["ACU-001", "Schroeder Reverb 合规修复", "P0", "全通滤波器 + 脉冲响应验证"],
        ["ACU-002", "RBJ Biquad EQ 替换", "P0", "RBJ biquad 替代 FFT sigmoid/Gaussian"],
        ["ACU-003", "HPSS 残差守恒", "P0", "保留 R 分量 + 能量守恒审计"],
        ["ACU-004", "7 频段默认启用", "P1", "增加 5-8 kHz Brilliance 区间"],
        ["ACU-005", "True Peak Limiter", "P1", "4x 过采样真峰值检测"],
        ["ACU-006", "Mel/Bark/ERB 感知尺度", "P1", "三种感知频率映射 + 频谱表示"],
        ["ACU-007", "心理声学掩蔽初版", "P2", "同时掩蔽 + 扩展函数最小模型"],
        ["ACU-008", "F0/Pitch Stability", "P2", "YIN 算法 + F0 稳定性度量"],
        ["ACU-009", "Chroma/Key/Harmony", "P2", "12 维 chroma + 调性检测"],
        ["ACU-010", "MRS 参考集鲁棒化", "P2", "分风格参考 + 鲁棒评分 + Bootstrap CI"],
    ]
    add_styled_table(doc, ["编号", "任务", "优先级", "核心"], aep_summary, [2, 5, 1.5, 7])

    add_heading(doc, "依赖关系", 2)
    doc.add_paragraph(
        "P0 三个任务（ACU-001/002/003）完全独立，可并行执行。"
        "P1 三个任务（ACU-004/005/006）完全独立。"
        "P2 中 ACU-007 依赖 ACU-006 的 Bark 尺度映射；其余 P2 任务独立。"
    )

    doc.add_page_break()

    # ── Chapter 4: Acceptance Matrix ──
    add_heading(doc, "第四部分：验收矩阵", 1)
    acceptance = [
        ["ACU-001", "全通级频谱平坦性", "< 0.1 dB", "白噪声频谱对比", "MUST"],
        ["ACU-001", "脉冲响应无离散回声", "峰值比 < 3:1", "单位脉冲峰值检测", "MUST"],
        ["ACU-001", "MRS texture 提升", "≥ 3 分中位数", "20 首配对 t 检验", "MUST"],
        ["ACU-002", "频率响应精度", "RMSE < 0.1 dB", "对数扫频 vs 理论", "MUST"],
        ["ACU-002", "零增益透明性", "RMSE < -96 dBFS", "零参数输出 vs 输入", "MUST"],
        ["ACU-002", "MRS 回归", "差值 < 2.0 分", "20 首 × 3 参数组", "MUST"],
        ["ACU-003", "能量守恒审计", "≤ 3σ safe", "conservation.py 审计", "MUST"],
        ["ACU-003", "MRS 不劣于当前", "p > 0.05", "20 首配对 t 检验", "MUST"],
        ["ACU-004", "单一定义来源", "bands.py 唯一来源", "代码审查", "MUST"],
        ["ACU-004", "频段能量守恒", "误差 < 0.01%", "50 首求和验证", "MUST"],
        ["ACU-005", "真峰值限幅", "真峰值 ≤ ceiling", "15 kHz @ -1 dBFS", "MUST"],
        ["ACU-005", "全频段安全", "≤ ceiling + 0.1 dB", "20 Hz-20 kHz 扫频", "MUST"],
        ["ACU-006", "感知频谱 shape", "维度正确", "与 librosa 对比", "MUST"],
        ["ACU-007", "掩蔽阈值定性", "峰值在 1 Bark", "1 kHz 纯音测试", "MUST"],
        ["ACU-008", "YIN 精度", "RMSE < 1 Hz", "与 librosa.yin 对比", "MUST"],
        ["ACU-009", "Chroma 归一化", "每帧和 = 1.0", "单元测试", "MUST"],
        ["ACU-010", "CI 宽度改善", "≥ 20% 更窄", "Bootstrap 1000 次", "MUST"],
    ]
    add_styled_table(doc, ["AEP", "验收项", "量化标准", "验证方法", "等级"], acceptance, [1.8, 4, 3.5, 4.5, 1.5])

    add_heading(doc, "全局验收门 (v0.4 GA Gate)", 2)
    gates = [
        ["GATE-01", "所有 MUST 验收项通过"],
        ["GATE-02", "所有 SHOULD 验收项有明确状态"],
        ["GATE-03", "ruff lint 零错误"],
        ["GATE-04", "pytest -m v01 全通过"],
        ["GATE-05", "pytest 全量通过"],
        ["GATE-06", "MRS 回归: 20 首基线差值 < 2.0 分"],
        ["GATE-07", "A_compliance ≥ 85 (vs 当前 65)"],
        ["GATE-08", "文档无 TODO/占位符"],
        ["GATE-09", "Founder/CTO 签核"],
    ]
    add_styled_table(doc, ["门编号", "条件"], gates, [2.5, 13])

    doc.add_page_break()

    # ── Chapter 5: Release Cadence ──
    add_heading(doc, "第五部分：研发节奏与冻结规则", 1)
    phases = [
        ["P0", "章程签核", "0.5 天", "Founder/CTO 审阅批准", "GATE-09 签核通过"],
        ["P1", "P0 修复", "2 天", "ACU-001/002/003", "3 个 P0 MUST 通过"],
        ["P2", "P1 补完", "5 天", "ACU-004/005/006", "3 个 P1 MUST 通过"],
        ["P3", "回归测试", "1 天", "全量 pytest + MRS 回归", "GATE-03~06 通过"],
        ["P4", "P2 判断", "1 天", "评估 P2 合入/推迟", "决定已记录"],
        ["P5", "封口归档", "1 天", "DOCX/PDF + git tag", "GATE-01~09 通过"],
    ]
    add_styled_table(doc, ["阶段", "名称", "耗时", "关键活动", "退出条件"], phases, [1.2, 2, 1.5, 5, 5])

    add_heading(doc, "冻结门禁", 2)
    freezes = [
        ["F-INPUT", "输入冻结", "P1 开始前", "范围边界 + AEP 任务 + 测试基线"],
        ["F-P0", "P0 实现冻结", "P1 完成时", "ACU-001/002/003 代码 + 验收证据"],
        ["F-P1", "P1 实现冻结", "P2 完成时", "ACU-004/005/006 代码"],
        ["F-MRS", "MRS 参考基线冻结", "P3 通过后", "20 首基线 MRS + 管线快照"],
        ["F-RELEASE", "发布冻结", "P5 完成时", "v0.4 全部代码/文档/基线 + git tag"],
    ]
    add_styled_table(doc, ["代号", "名称", "触发时间", "冻结内容"], freezes, [2, 3, 2.5, 8])

    doc.add_page_break()

    # ── Chapter 6: Risk Register ──
    add_heading(doc, "第六部分：风险登记表", 1)
    risks = [
        ["R-006", "回归测试覆盖不足", "3", "5", "15", "阻断", "先写测试，再改代码"],
        ["R-001", "RBJ EQ 替换 MRS 回归失败", "3", "4", "12", "高", "提前确认差异量级"],
        ["R-007", "P0 修复间意外交互", "3", "4", "12", "高", "端到端组合测试"],
        ["R-008", "P2 范围蔓延延期", "4", "3", "12", "高", "硬截止 Day 9 15:00"],
        ["R-004", "频段变更历史数据不可比", "5", "2", "10", "高", "保留 6 频段兼容"],
        ["R-009", "新 P0 缺陷发现", "2", "5", "10", "高", "评估后可能拆分版本"],
        ["R-003", "HPSS 残差引入噪声", "3", "3", "9", "中", "R_process 策略"],
        ["R-002", "全通级低频相位抵消", "2", "4", "8", "中", "降低 g1 参数"],
        ["R-005", "真峰值过采样延迟", "3", "2", "6", "中", "高效 polyphase 实现"],
        ["R-010", "签核延迟", "3", "2", "6", "中", "异步审阅 + 假定批准"],
    ]
    add_styled_table(doc, ["编号", "风险", "P", "I", "P×I", "等级", "缓解措施"],
                    risks, [1.5, 4.5, 1, 1, 1.2, 1.2, 5])

    doc.add_page_break()

    # ── Chapter 7: Formula System ──
    add_heading(doc, "第七部分：公式与指标体系", 1)

    add_heading(doc, "v0.4 完成度公式", 2)
    p = doc.add_paragraph()
    run = p.add_run("C_v04 = S_scope × A_AEP × Q_accept × F_freeze / R_drift")
    run.bold = True
    run.font.size = Pt(11)

    formula_vars = [
        ["S_scope", "范围清晰度", "0.95 → 0.98", "In/Out Scope 完整性"],
        ["A_AEP", "AEP 任务完整度", "1.00 → 1.00", "10 张任务卡定义完整"],
        ["Q_accept", "验收标准质量", "0.95 → 1.00", "MUST/SHOULD/MAY 有量化标准"],
        ["F_freeze", "版本冻结能力", "1.00 → 1.00", "5 个冻结门禁清晰可执行"],
        ["R_drift", "范围漂移风险", "0.15 → <0.10", "范围变更概率 × 影响"],
    ]
    add_styled_table(doc, ["变量", "名称", "当前→目标", "测量方法"], formula_vars, [2, 3, 3, 7.5])

    add_heading(doc, "v0.4 健康度仪表盘", 2)
    dashboard = [
        ["A_compliance", "65/100", "88/100", "+23", "ACU-001/002/003/005"],
        ["P_align", "0/100", "3/100", "+3", "ACU-006/007/008/009"],
        ["C_loop", "6.0%", "14.4%", "+8.4pp", "全维度改善"],
        ["C_v04", "—", "9.8/10", "—", "章程执行"],
        ["MRS 回归稳定性", "—", "< 2.0 分", "—", "ACU-002 验证"],
    ]
    add_styled_table(doc, ["指标", "当前", "目标", "变化", "驱动因素"], dashboard, [3, 2.5, 2.5, 2.5, 5])

    doc.add_page_break()

    # ── Chapter 8: DeepSeek Execution Prompt ──
    add_heading(doc, "第八部分：DeepSeek 执行提示词（摘要）", 1)
    doc.add_paragraph(
        "本部分为 DeepSeek (AI Worker Layer) 执行 AEP-ACU 任务时的核心约束摘要。"
        "完整版见 markdown/07_deepseek_execution_prompt.md。"
    )

    add_heading(doc, "执行硬约束", 2)
    constraints = [
        "每个输出必须落到：变量、测量值、统计检验、验收证据或冻结判定。",
        "不允许只写描述性段落。无数字/测试结果/对比数据/验证结论的段落须删除。",
        "不确定的内容必须标记为 [理论假设] 或 [待实验验证]。",
        "每个 AEP 执行须产出四类产物：实验设计、原始测量数据、统计分析、验收判定。",
        "不允许跳过冻结标准中的任何 MUST 项。",
    ]
    for c in constraints:
        doc.add_paragraph(c, style="List Bullet")

    add_heading(doc, "禁止行为", 2)
    forbidden = [
        "不写没有数字的段落",
        "不跳过冻结标准的 MUST 项",
        "不修改验收标准以适应实测结果",
        "不在实验数据不完整时声称「预期通过」",
        "不将一个 AEP 的结论建立在另一个未完成 AEP 的假设之上",
        "不将公开标准算法描述为 Moodify 的发明",
    ]
    for f in forbidden:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_page_break()

    # ── Sign-off page ──
    add_heading(doc, "签核页", 1)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    signoff = [
        ["Founder / CTO", "________________", "____/____/________"],
        ["Claude A (交接官)", "________________", "____/____/________"],
    ]
    add_styled_table(doc, ["角色", "签名", "日期"], signoff, [5, 5, 5])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "签核后 DOC-MFY-003 进入 Phase 1 (P0 修复)。"
        "v0.4 封口目标日期: 2026-07-15。"
    )
    run.italic = True

    add_page_number(doc)
    path = OUT_DIR / "DOC-MFY-003_项目章程_合订版.docx"
    doc.save(str(path))
    print(f"  -> {path}")


# ═══════════════════════════════════════════════════════════════
# 2. 需求说明书
# ═══════════════════════════════════════════════════════════════
def build_requirements():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Cover
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DOC-MFY-003\n需求说明书")
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(*_hex_to_rgb(HEADER_COLOR))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Moodify Acoustic Compliance Upgrade v0.4\n2026-07-02")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(*_hex_to_rgb(ACCENT_COLOR))

    doc.add_page_break()

    add_heading(doc, "1. 需求概述", 1)
    add_meta_table(doc)

    add_heading(doc, "1.1 项目目标", 2)
    doc.add_paragraph(
        "将 Moodify 声学合规度 A_compliance 从 65/100 提升至 88/100。"
        "修复 DOC-MFY-002 审计发现的 P0 级声学理论与实现缺口，"
        "建立感知声学、音乐智能和 MRS 鲁棒化的基础设施入口。"
    )

    add_heading(doc, "1.2 核心约束", 2)
    constraints = [
        "不做大模型训练 — 硬件不足（RTX 4060 仅 8GB VRAM）",
        "不做 UI/UX 改版 — v0.4 是声学引擎升级",
        "不引入新的外部 API 依赖",
        "不新增 CLI 命令或 API 端点 — 修复对现有接口透明",
        "不做大规模架构重构 — 修复在现有模块边界内完成",
    ]
    for c in constraints:
        doc.add_paragraph(c, style="List Bullet")

    add_heading(doc, "2. 功能需求", 1)

    # ── P0 ──
    add_heading(doc, "2.1 P0 — 阻塞发布", 2)

    add_heading(doc, "FR-001: Schroeder Reverb 合规修复 (ACU-001)", 3)
    req1 = [
        ("输入", "processing/operators.py:262-275 的当前 _schroeder_reverb()"),
        ("需求", "在梳状滤波器循环后添加 2 个全通滤波器级（全通公式: y[n] = -g*x[n] + x[n-K] + g*y[n-K]）"),
        ("参数", "K1 = int(sr * 0.005), g1 = 0.7; K2 = int(sr * 0.0017), g2 = 0.7"),
        ("验收", "全通级幅频响应平坦度 < 0.1 dB; 脉冲响应 t > 50 ms 内无可辨识离散回声; MRS texture 提升 ≥ 3 分"),
    ]
    for label, text in req1:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：")
        run.bold = True
        p.add_run(text)

    add_heading(doc, "FR-002: RBJ Biquad EQ 替换 (ACU-002)", 3)
    req2 = [
        ("输入", "processing/operators.py:116-138 的当前 FFT sigmoid/Gaussian EQ"),
        ("需求", "用 RBJ biquad 标准滤波器（low_shelf / high_shelf / peaking）替换 FFT EQ"),
        ("方案 A", "废弃 operators.py 的 FFT EQ，使用 pedalboard 的 PeakFilter / LowShelfFilter / HighShelfFilter"),
        ("方案 B", "从 craft_processes.py 提取 RBJ biquad 实现为独立模块 processing/rbj_eq.py"),
        ("验收", "扫频 RMSE < 0.1 dB; 零增益输出 vs 输入 RMSE < -96 dBFS; MRS 回归差值 < 2.0 分"),
    ]
    for label, text in req2:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：")
        run.bold = True
        p.add_run(text)

    add_heading(doc, "FR-003: HPSS 残差守恒 (ACU-003)", 3)
    req3 = [
        ("输入", "processing/spectral_chain.py 的当前 HPSS 处理流程（丢弃 R 分量）"),
        ("需求", "保留 Harmonic、Percussive、Residual 三个分量；审计残差能量；选择最优策略（R_add_back 或 R_process）"),
        ("验收", "能量审计 |ΔL_residual| ≤ 3σ; MRS 不劣于当前; 无新增可听伪影"),
    ]
    for label, text in req3:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：")
        run.bold = True
        p.add_run(text)

    # ── P1 ──
    add_heading(doc, "2.2 P1 — 必须完成", 2)

    add_heading(doc, "FR-004: 7 频段默认启用 (ACU-004)", 3)
    req4 = [
        ("输入", "bands.py 的当前 6 频段定义"),
        ("需求", "扩展为 7 频段：Sub(20-60) / Bass(60-250) / Low-Mid(250-500) / Mid(500-2000) / Presence(2000-5000) / Brilliance(5000-8000) / Air(8000-16000)"),
        ("验收", "bands.py 是唯一频段定义来源; 7 频段能量总和 = 全频段能量（误差 < 0.01%）; Brilliance 区间 AI vs 真实 d > 0.3"),
    ]
    for label, text in req4:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：")
        run.bold = True
        p.add_run(text)

    add_heading(doc, "FR-005: True Peak Limiter (ACU-005)", 3)
    req5 = [
        ("输入", "processing/operators.py:309-343 的当前采样峰值砖墙限幅器"),
        ("需求", "升级为 4x 过采样真峰值砖墙限幅器（参照 ITU-R BS.1771-1）+ 添加 1 ms 非零 attack"),
        ("验收", "15 kHz @ -1 dBFS → 真峰值 ≤ ceiling; 全频段真峰值 ≤ ceiling + 0.1 dB; 低频 THD < 0.5%"),
    ]
    for label, text in req5:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：")
        run.bold = True
        p.add_run(text)

    add_heading(doc, "FR-006: Mel/Bark/ERB 感知尺度 (ACU-006)", 3)
    req6 = [
        ("输入", "diagnosis/metrics.py 的当前线性 FFT 频谱分析器"),
        ("需求", "实现 bark_spectrogram()、mel_spectrogram()、erb_spectrogram() 三个感知频谱函数；SpectrumAnalyzer 增加 scale 参数"),
        ("验收", "三种感知频谱 shape 正确; 60-250 Hz 在 Bark 尺度上由 2-3 个 band 表示; 计算时间 < 2x 线性 FFT"),
    ]
    for label, text in req6:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：")
        run.bold = True
        p.add_run(text)

    # ── P2 ──
    add_heading(doc, "2.3 P2 — 条件合入", 2)

    add_heading(doc, "FR-007: 心理声学掩蔽初版 (ACU-007)", 3)
    p = doc.add_paragraph()
    run = p.add_run("需求：")
    run.bold = True
    p.add_run("Bark 尺度映射 + 扩展函数 + 掩蔽阈值计算。依赖 ACU-006 先完成。")

    add_heading(doc, "FR-008: F0/Pitch Stability (ACU-008)", 3)
    p = doc.add_paragraph()
    run = p.add_run("需求：")
    run.bold = True
    p.add_run("YIN 算法 F0 检测 + 音高稳定性度量。AI vs 真实录音 F0 稳定性 d > 0.5。")

    add_heading(doc, "FR-009: Chroma/Key/Harmony (ACU-009)", 3)
    p = doc.add_paragraph()
    run = p.add_run("需求：")
    run.bold = True
    p.add_run("12 维 chroma + Krumhansl-Schmuckler 调性检测 + 和声稳定性。调性准确率 > 80%。")

    add_heading(doc, "FR-010: MRS 参考集鲁棒化 (ACU-010)", 3)
    p = doc.add_paragraph()
    run = p.add_run("需求：")
    run.bold = True
    p.add_run("分风格参考分布 + MAD 鲁棒评分 + Bootstrap 95% CI。分风格 CI 比混合 CI 窄 ≥ 20%。")

    doc.add_page_break()

    add_heading(doc, "3. 非功能需求", 1)
    nfr = [
        ["NFR-01", "代码质量", "ruff lint 零错误; pytest 全量通过; 覆盖率 ≥ 70%"],
        ["NFR-02", "性能", "EQ 替换后处理延迟增加 < 10%; True Peak 延迟增加 < 10 ms; 感知频谱 < 2x 线性 FFT"],
        ["NFR-03", "兼容性", "输入/输出格式不变; 保留 6 频段向后兼容; 诊断报告标注频段版本"],
        ["NFR-04", "安全性", "不引入新的外部 API 依赖; 所有第三方库固定版本"],
        ["NFR-05", "可追溯性", "每个 AEP 的验收证据保存在 executions/ACU-NNN/; MRS 分数可追溯到基线集"],
    ]
    add_styled_table(doc, ["编号", "类别", "要求"], nfr, [2, 2.5, 11])

    add_heading(doc, "4. 测试基线集要求", 1)
    baseline = [
        ["数量", "20 首（固定，不可替换）"],
        ["来源", "AI 生成 ≥ 12 首 + 真实录音 ≥ 8 首"],
        ["风格", "Classical / Jazz / Rock / Pop / Electronic / Acoustic 各 ≥ 2 首"],
        ["格式", "WAV 44.1 kHz 16-bit stereo, 30-180 秒"],
        ["路径", "tests/data/baseline/v04/"],
        ["冻结", "M2 里程碑前冻结，不可替换或添加"],
    ]
    add_styled_table(doc, ["要求", "规格"], baseline, [2.5, 13])

    add_heading(doc, "5. 签核", 1)
    doc.add_paragraph()
    sig = [
        ["Founder / CTO", "________________", "____/____/________"],
    ]
    add_styled_table(doc, ["角色", "签名", "日期"], sig, [5, 5, 5])

    add_page_number(doc)
    path = OUT_DIR / "DOC-MFY-003_需求说明书.docx"
    doc.save(str(path))
    print(f"  -> {path}")


# ═══════════════════════════════════════════════════════════════
# 3. 验收标准与检查表
# ═══════════════════════════════════════════════════════════════
def build_acceptance():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Cover
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DOC-MFY-003\n验收标准与检查表")
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(*_hex_to_rgb(HEADER_COLOR))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Moodify Acoustic Compliance Upgrade v0.4\n2026-07-02")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(*_hex_to_rgb(ACCENT_COLOR))

    doc.add_page_break()

    add_heading(doc, "1. 验收体系说明", 1)
    doc.add_paragraph(
        "本检查表基于 DOC-MFY-003 验收矩阵（03_acceptance_matrix.md），"
        "将所有 AEP 的验收项转换为可逐项打勾的核查表。"
        "验收等级：MUST = 不通过则 v0.4 不发布 | SHOULD = 需要合理理由才能跳过 | MAY = 条件允许则合入。"
    )

    add_heading(doc, "2. 逐 AEP 验收检查", 1)

    aep_checks = [
        ("AEP-ACU-001: Schroeder Reverb 合规修复 [P0]", [
            ("□", "MUST", "全通级幅频响应平坦度 < 0.1 dB (白噪声测试)"),
            ("□", "MUST", "脉冲响应 t > 50 ms 内无可辨识离散回声 (峰值比 < 3:1)"),
            ("□", "MUST", "MRS texture 组件提升 ≥ 3 分 (20 首测试音频中位数)"),
            ("□", "MUST", "MRS space 组件不降低 (配对 t 检验 p > 0.05)"),
            ("□", "MUST", "ruff lint 通过"),
            ("□", "MUST", "pytest -m v01 通过"),
        ]),
        ("AEP-ACU-002: RBJ Biquad EQ 替换 [P0]", [
            ("□", "MUST", "扫频测试：20 Hz – 20 kHz RBJ vs 理论 RMSE < 0.1 dB"),
            ("□", "MUST", "零增益测试：所有 EQ 参数=0 时，输出 vs 输入 RMSE < -96 dBFS"),
            ("□", "MUST", "MRS 回归：20 首测试音频前后差值变化 < 2.0 分"),
            ("□", "SHOULD", "处理延迟不增加 > 10%"),
            ("□", "MUST", "ruff lint 通过"),
            ("□", "MUST", "pytest -m v01 通过 + pytest 全量通过"),
        ]),
        ("AEP-ACU-003: HPSS 残差守恒 [P0]", [
            ("□", "MUST", "能量审计: |ΔL_residual| ≤ 3σ (safe 级别)"),
            ("□", "MUST", "MRS 不低于当前 (配对 t 检验 p > 0.05)"),
            ("□", "SHOULD", "无新增可听伪影 (非正式听感检查 + MRS artifact 组件不降低)"),
            ("□", "MUST", "ruff lint 通过"),
            ("□", "MUST", "pytest -m v01 通过"),
        ]),
        ("AEP-ACU-004: 7 频段默认启用 [P1]", [
            ("□", "MUST", "bands.py 是唯一的频段定义来源 (消除 DEF-006 不一致)"),
            ("□", "MUST", "7 频段能量总和 = 全频段能量 (误差 < 0.01%)"),
            ("□", "SHOULD", "5-8 kHz Brilliance 区间 AI vs 真实 d > 0.3"),
            ("□", "MUST", "ruff lint 通过"),
            ("□", "MUST", "pytest 通过"),
        ]),
        ("AEP-ACU-005: True Peak Limiter [P1]", [
            ("□", "MUST", "15 kHz @ -1 dBFS 输入 → 真峰值输出 ≤ ceiling"),
            ("□", "MUST", "全频段扫频：真峰值始终 ≤ ceiling + 0.1 dB"),
            ("□", "SHOULD", "低频 (< 100 Hz) 限幅后 THD < 0.5%"),
            ("□", "MUST", "处理延迟增加 < 10 ms"),
            ("□", "MUST", "MRS 回归通过 (20 首差值 < 2.0 分)"),
            ("□", "MUST", "ruff lint + pytest 通过"),
        ]),
        ("AEP-ACU-006: Mel/Bark/ERB 感知尺度 [P1]", [
            ("□", "MUST", "Bark/Mel/ERB 三种频谱输出 shape 正确"),
            ("□", "SHOULD", "低频区域 (60-250 Hz) 在 Bark 尺度上由 ≤ 4 个 band 表示"),
            ("□", "SHOULD", "三种感知频谱的计算时间 < 2x 线性 FFT"),
            ("□", "MUST", "ruff lint + pytest 通过"),
        ]),
        ("AEP-ACU-007: 心理声学掩蔽初版 [P2]", [
            ("□", "MUST", "1 kHz 纯音掩蔽阈值峰值在 1 Bark 附近 (定性验证)"),
            ("□", "SHOULD", "扩展函数不对称性验证: 低频→高频 > 高频→低频"),
            ("□", "SHOULD", "3 分钟音频掩蔽计算时间 < 10 秒"),
            ("□", "MUST", "ruff lint + pytest 通过"),
        ]),
        ("AEP-ACU-008: F0/Pitch Stability [P2]", [
            ("□", "MUST", "YIN 输出与 librosa.yin 一致 (RMSE < 1 Hz for f0 < 1000 Hz)"),
            ("□", "SHOULD", "AI 音频 F0 稳定性中位数 vs 真实录音: d > 0.5"),
            ("□", "SHOULD", "3 分钟音频 F0 计算时间 < 5 秒"),
            ("□", "MUST", "ruff lint + pytest 通过"),
        ]),
        ("AEP-ACU-009: Chroma/Key/Harmony [P2]", [
            ("□", "MUST", "Chroma 向量和为 1 (归一化验证, tol 1e-6)"),
            ("□", "SHOULD", "调性检测在已知调性测试音频上准确率 > 80%"),
            ("□", "SHOULD", "AI 音频和声稳定性中位数 vs 真实录音: d > 0.3"),
            ("□", "MUST", "ruff lint + pytest 通过"),
        ]),
        ("AEP-ACU-010: MRS 参考集鲁棒化 [P2]", [
            ("□", "MUST", "分风格 MRS 95% CI 宽度比混合 MRS 窄 ≥ 20%"),
            ("□", "SHOULD", "MAD-based 评分在 contaminated sample 上偏差 < SD-based 评分"),
            ("□", "SHOULD", "Bootstrap CI 覆盖真实 MRS 值的比例 ≥ 93%"),
            ("□", "MUST", "ruff lint + pytest 通过"),
        ]),
    ]

    for aep_name, checks in aep_checks:
        add_heading(doc, aep_name, 3)
        for checkbox, level, desc in checks:
            p = doc.add_paragraph()
            run = p.add_run(f"{checkbox} [{level}] ")
            run.bold = True
            p.add_run(desc)

    doc.add_page_break()

    add_heading(doc, "3. 全局验收门检查表", 1)
    gates = [
        ("□", "GATE-01", "所有 MUST 验收项通过"),
        ("□", "GATE-02", "所有 SHOULD 验收项有明确状态（通过/跳过+理由/失败）"),
        ("□", "GATE-03", "ruff lint 零错误 (ruff check .)"),
        ("□", "GATE-04", "pytest -m v01 全通过"),
        ("□", "GATE-05", "pytest 全量通过"),
        ("□", "GATE-06", "MRS 回归: 20 首基线音频处理前后 MRS 差值变化 < 2.0 分"),
        ("□", "GATE-07", "A_compliance ≥ 85 (vs 当前 65)"),
        ("□", "GATE-08", "DOC-MFY-003 全部 Markdown 文件无 TODO/占位符"),
        ("□", "GATE-09", "Founder/CTO 签核"),
    ]
    for checkbox, gid, desc in gates:
        p = doc.add_paragraph()
        run = p.add_run(f"{checkbox} {gid}: ")
        run.bold = True
        p.add_run(desc)

    add_heading(doc, "4. 测试基线集检查", 1)
    baseline_checks = [
        ("□", "20 首固定基线音频已选定并冻结"),
        ("□", "AI 生成至少 12 首 (Suno v3/v4, Udio v1)"),
        ("□", "真实录音至少 8 首"),
        ("□", "6 种风格各 ≥ 2 首 (Classical/Jazz/Rock/Pop/Electronic/Acoustic)"),
        ("□", "格式: WAV 44.1 kHz 16-bit stereo, 30-180 秒"),
        ("□", "存储路径: tests/data/baseline/v04/"),
        ("□", "M2 里程碑前冻结 —— 不可替换或添加"),
    ]
    for checkbox, desc in baseline_checks:
        p = doc.add_paragraph()
        run = p.add_run(f"{checkbox} ")
        run.bold = True
        p.add_run(desc)

    add_heading(doc, "5. 文档完整性检查", 1)
    doc_checks = [
        ("□", "00_project_charter.md — 项目章程正文"),
        ("□", "01_scope_boundary.md — 范围边界与不做事项"),
        ("□", "02_aep_acu_list.md — AEP-ACU-001~010 任务卡"),
        ("□", "03_acceptance_matrix.md — 验收矩阵"),
        ("□", "04_release_cadence.md — 研发节奏与冻结规则"),
        ("□", "05_risk_register.md — 风险登记表"),
        ("□", "06_formula_system.md — 公式与指标体系"),
        ("□", "07_deepseek_execution_prompt.md — DeepSeek 执行提示词"),
        ("□", "metadata.json — 机器可读元数据"),
        ("□", "README.txt — 项目说明"),
        ("□", "references/reference_standards.md — 参考标准"),
        ("□", "assets/ — 图表资产 (scope/priority/cadence)"),
        ("□", "docx/ — Word 交付物 (3 files)"),
        ("□", "pdf/ — PDF 交付物 (3 files)"),
    ]
    for checkbox, desc in doc_checks:
        p = doc.add_paragraph()
        run = p.add_run(f"{checkbox} ")
        run.bold = True
        p.add_run(desc)

    add_heading(doc, "6. 签核", 1)
    doc.add_paragraph()
    sig = [
        ["Founder / CTO", "________________", "____/____/________"],
        ["Claude A (交接官)", "________________", "____/____/________"],
    ]
    add_styled_table(doc, ["角色", "签名", "日期"], sig, [5, 5, 5])

    add_page_number(doc)
    path = OUT_DIR / "DOC-MFY-003_验收标准与检查表.docx"
    doc.save(str(path))
    print(f"  -> {path}")


# ═══════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Generating DOC-MFY-003 DOCX deliverables...")
    build_charter()
    build_requirements()
    build_acceptance()
    print("Done.")
